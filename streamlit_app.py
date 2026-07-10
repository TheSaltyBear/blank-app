import streamlit as st
import pandas as pd
import numpy as np
import requests
import re
import os
import matplotlib.pyplot as plt
from datetime import datetime, timedelta

API_KEY = os.environ.get("OPENAI_API_KEY","")

st.set_page_config(page_title="VA Inventory Intelligence",page_icon="🏥",layout="wide")

def parse_usage(text):
    recs,cur,monthly=[],{},{}
    for L in text.splitlines():
        if L.startswith("GROUP:"):
            cur,monthly={},{}
            m=re.search(r"\(#(\d+)\)",L)
            cur["item_id"]=m and int(m.group(1))
        m=re.match(r"^(.+?)\s+(\d+)\s+([\d/]+)\s+([\d.]+)\s+([\d.]+)\s+(\d+)",L)
        if m and "description" not in cur:
            cur.update({"description":m[1].strip(),"on_hand":int(m[2]),"unit":m[3],"last_cost":float(m[4]),"avg_cost":float(m[5])})
        m=re.search(r"NORM:\s*(\d+)\s+REORD:\s*(\d+)\s+OPT:\s*(\d*)\s+EMER:\s*(\d+)",L)
        if m:
            cur.update({"norm_stk":int(m[1]),"reorder_pt":int(m[2]),"opt_lvl":int(m[3] or 0),"emerg_lvl":int(m[4])})
        months=re.findall(r"([A-Z]{3}\d{2})\s+(\d+)\s+([\d.]+)",L)
        for mo,qty,cost in months:
            monthly[mo+"_qty"]=int(qty)
            monthly[mo+"_cost"]=float(cost)
        if L.strip().startswith("CUMULATIVE TOTAL"):
            m=re.search(r"CUMULATIVE TOTAL\s+(\d+)\s+([\d.]+)",L)
            if m:
                cur["cumulative_qty"]=int(m[1])
                cur["cumulative_cost"]=float(m[2])
                cur.update(monthly)
                recs.append(cur.copy())
                cur,monthly={},{}
    return pd.DataFrame(recs) if recs else pd.DataFrame()

def parse_comp(text):
    recs,cur=[],{}
    for L in text.splitlines():
        m=re.search(r"\[#(\d+)\]",L)
        if m:
            if cur: recs.append(cur.copy())
            cur={"item_id":int(m.group(1)),"description":L.split("[")[0].strip()}
        if not cur: continue
        m=re.search(r"ON-DEMAND:\s*(\w)\s+SUBACCOUNT:\s*(\d+)",L)
        if m: cur["on_demand"]=m[1]; cur["subacct"]=int(m[2])
        m=re.search(r"QTY ON HAND:\s*(\d+)\s+DUE-IN:\s*(\d+)\s+DUE-OUT:\s*(\d+)",L)
        if m: cur["qty_on_hand"]=int(m[1]); cur["due_in"]=int(m[2]); cur["due_out"]=int(m[3])
        m=re.search(r"NORM STK LVL:\s*(\d+)\s+REORDER PT:\s*(\d+)",L)
        if m: cur["norm_stk"]=int(m[1]); cur["reorder_pt"]=int(m[2])
        m=re.search(r"EMERGENCY LVL:\s*(\d+)",L)
        if m: cur["emerg_lvl"]=int(m[1])
        m=re.search(r"LAST COST:\s*([\d.]+)\s+AVERAGE COST:\s*([\d.]+)",L)
        if m: cur["last_cost"]=float(m[1]); cur["avg_cost"]=float(m[2])
        m=re.search(r"LAST REC'D:\s*([A-Z]+\s+\d+,\s*\d+)",L)
        if m: cur["last_rcvd"]=m[1]
        m=re.search(r"MAIN STORAGE LOC:\s*(\S+)",L)
        if m: cur["location"]=m[1]
    if cur: recs.append(cur.copy())
    return pd.DataFrame(recs) if recs else pd.DataFrame()

def call_llm(prompt):
    try:
        resp=requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization":f"Bearer {API_KEY}","Content-Type":"application/json"},
            json={"model":"gpt-4o-mini","messages":[
                {"role":"system","content":"You are a senior VA hospital inventory expert with 20 years experience. You understand VA procurement rules, PAR levels, mandatory sources, lead times, and patient care impact. Give direct, actionable advice."},
                {"role":"user","content":prompt}
            ],"temperature":0.2}
        )
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"LLM error: {e}"

def risk_score(row):
    score=0
    if row.get("on_hand",0)<=row.get("emerg_lvl",0): score+=40
    elif row.get("on_hand",0)<=row.get("reorder_pt",0): score+=25
    if row.get("due_in",0)==0: score+=20
    try:
        last=datetime.strptime(row.get("last_rcvd",""),"%b %d, %Y")
        days=(datetime.today()-last).days
        if days>90: score+=20
        elif days>30: score+=10
    except: score+=10
    return min(score,100)

def simple_forecast(row):
    months=["APR26","MAY26","JUN26","JUL26"]
    qtys=[float(row[mo+"_qty"]) for mo in months if mo+"_qty" in row and pd.notna(row[mo+"_qty"])]
    if len(qtys)<2: return None,None
    avg=np.mean(qtys)
    trend=(qtys[-1]-qtys[0])/len(qtys)
    return avg,[max(0,avg+trend*i) for i in range(1,4)]

st.title("🏥 VA Inventory Intelligence")
st.caption("Phoenix VA Hospital — Powered by AI")

with st.sidebar:
    st.header("📂 Upload Morning Reports")
    usage_file=st.file_uploader("Usage Demand Report",type=["txt","docx"])
    comp_file=st.file_uploader("Comp Item List",type=["txt","docx"])
    run_btn=st.button("🚀 Run Analysis",use_container_width=True)
    st.markdown("---")
    st.caption("Upload both files then tap Run Analysis")

if run_btn and usage_file and comp_file:
    if usage_file.name.endswith(".docx"):
        from docx import Document
        import io
        doc=Document(io.BytesIO(usage_file.read()))
        utext="\n".join([p.text for p in doc.paragraphs])
    else:
        utext=usage_file.read().decode("utf-8",errors="ignore")
    if comp_file.name.endswith(".docx"):
        from docx import Document
        import io
        doc=Document(io.BytesIO(comp_file.read()))
        ctext="\n".join([p.text for p in doc.paragraphs])
    else:
        ctext=comp_file.read().decode("utf-8",errors="ignore")

    with st.spinner("Parsing reports..."):
        u=parse_usage(utext)
        c=parse_comp(ctext)

    if u.empty or c.empty:
        st.error("Could not parse files. Please check your uploads.")
        st.stop()

    if "description" in u.columns and "description" in c.columns:
        c_cols=[col for col in ["description","qty_on_hand","due_in","due_out","last_rcvd","location","avg_cost"] if col in c.columns]
        c_dedup=c[c_cols].drop_duplicates("description")
        df=u.merge(c_dedup,on="description",how="left")
    else:
        df=u.copy()

    df["qty_on_hand"]=df.get("qty_on_hand",df.get("on_hand",0))
    df["due_in"]=df.get("due_in",pd.Series([0]*len(df))).fillna(0)
    df["risk_score"]=df.apply(risk_score,axis=1)
    df["status"]=pd.cut(df["risk_score"],bins=[-1,30,60,100],labels=["✅ OK","⚠️ Watch","🚨 Critical"])

    st.markdown("## 📊 Today's Snapshot")
    k1,k2,k3,k4,k5=st.columns(5)
    k1.metric("Total Items",len(df))
    k2.metric("🚨 Critical",len(df[df.status=="🚨 Critical"]))
    k3.metric("⚠️ Watch",len(df[df.status=="⚠️ Watch"]))
    k4.metric("✅ OK",len(df[df.status=="✅ OK"]))
    k5.metric("Total Inventory $",f"${(df['qty_on_hand']*df['last_cost']).sum():,.2f}")
    st.markdown("---")

    st.markdown("## 🚨 Critical Items")
    crit=df[df.status=="🚨 Critical"].sort_values("risk_score",ascending=False)
    if len(crit):
        show_cols=[c for c in ["description","qty_on_hand","reorder_pt","emerg_lvl","due_in","last_cost","risk_score","location"] if c in crit.columns]
        st.dataframe(crit[show_cols],use_container_width=True)
    else:
        st.success("No critical items today!")

    st.markdown("## ⚠️ Watch Items")
    watch=df[df.status=="⚠️ Watch"].sort_values("risk_score",ascending=False)
    if len(watch):
        show_cols=[c for c in ["description","qty_on_hand","reorder_pt","due_in","last_cost","risk_score","location"] if c in watch.columns]
        st.dataframe(watch[show_cols],use_container_width=True)

    st.markdown("## 📈 30-Day Demand Forecast")
    top10=df.nlargest(10,"risk_score")
    fig,ax=plt.subplots(figsize=(10,4))
    plotted=0
    for _,row in top10.iterrows():
        avg,fc=simple_forecast(row)
        if fc:
            future=[datetime.today()+timedelta(days=30*i) for i in range(1,4)]
            ax.plot(future,fc,marker="o",label=str(row["description"])[:20])
            plotted+=1
    if plotted:
        ax.set_title("Forecasted Demand — Next 3 Months")
        ax.set_ylabel("Units")
        ax.legend(fontsize=7,loc="upper left")
        plt.tight_layout()
        st.pyplot(fig)
    else:
        st.info("Not enough monthly data for forecast yet.")

    st.markdown("## 💰 Cost Analysis")
    c1,c2=st.columns(2)
    with c1:
        st.markdown("**Top 10 Items by Inventory Value**")
        df["inv_value"]=df["qty_on_hand"]*df["last_cost"]
        top_val=df.nlargest(10,"inv_value")[["description","qty_on_hand","last_cost","inv_value"]]
        st.dataframe(top_val,use_container_width=True)
    with c2:
        st.markdown("**Monthly Spend Trend**")
        month_keys=["APR26","MAY26","JUN26","JUL26"]
        spend=[df[mo+"_cost"].sum() if mo+"_cost" in df.columns else 0 for mo in month_keys]
        fig2,ax2=plt.subplots(figsize=(5,3))
        ax2.bar(month_keys,spend,color="#1f77b4")
        ax2.set_title("Total Monthly Spend")
        ax2.set_ylabel("$")
        plt.tight_layout()
        st.pyplot(fig2)

    st.markdown("## 🐢 Slow Movers / Dead Stock")
    slow_keys=[mo+"_qty" for mo in month_keys if mo+"_qty" in df.columns]
    if slow_keys:
        df["total_usage"]=df[slow_keys].fillna(0).sum(axis=1)
        slow=df[(df["total_usage"]==0)&(df["qty_on_hand"]>df["norm_stk"])][["description","qty_on_hand","norm_stk","last_cost"]]
        if len(slow):
            slow["excess_value"]=(slow["qty_on_hand"]-slow["norm_stk"])*slow["last_cost"]
            st.dataframe(slow.sort_values("excess_value",ascending=False),use_container_width=True)
            st.warning(f"💸 Excess dead stock value: ${slow['excess_value'].sum():,.2f}")
        else:
            st.success("No dead stock detected!")

    st.markdown("## 🤖 AI Recommendations")
    with st.spinner("Asking the AI..."):
        top5=df.nlargest(5,"risk_score")
        bullets="\n".join(
            f"• {r.description}: on_hand={r.qty_on_hand}, reorder_pt={r.get('reorder_pt',0)}, emerg_lvl={r.get('emerg_lvl',0)}, due_in={r.get('due_in',0)}, last_cost=${r.last_cost}, risk_score={r.risk_score}/100"
            for _,r in top5.iterrows()
        )
        prompt=f"""
Today is {datetime.today().strftime('%B %d, %Y')}.
You are advising the inventory manager at Phoenix VA Hospital.

TOP 5 AT-RISK ITEMS:
{bullets}

Please provide:
1. IMMEDIATE ACTIONS — what needs to happen today
2. ORDER RECOMMENDATIONS — what to order, how much, from which source if known
3. SAFETY STOCK ADJUSTMENTS — any PAR level changes you recommend
4. DEAD STOCK ACTIONS — what to do with excess inventory
5. OVERALL RISK SUMMARY — one paragraph plain English summary

Be specific, direct, and actionable. This is a VA hospital — patient care depends on it.
"""
        advice=call_llm(prompt)
    st.markdown(advice)

    with st.expander("📋 View Full Inventory Table"):
        st.dataframe(df.sort_values("risk_score",ascending=False),use_container_width=True)

    st.success("✅ Analysis complete!")

else:
    st.info("👈 Upload your Usage and Comp reports in the sidebar, then tap Run Analysis.")
